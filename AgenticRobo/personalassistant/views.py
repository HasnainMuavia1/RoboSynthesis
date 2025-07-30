from django.shortcuts import render, redirect
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib import messages
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, StreamingHttpResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from .watson_services import text_to_speech, speech_to_text
from django.http.response import HttpResponse
from django.views.decorators.http import require_POST
from .models import SubjectContext
import json
import os
import requests
import uuid
import time
import base64
import io
import fitz  # PyMuPDF for PDF processing
import tempfile
import pandas as pd
from PIL import Image
import docx  # python-docx for DOCX processing

from docx import Document
from dotenv import load_dotenv
from langchain_tavily import TavilySearch
from langchain.memory import ConversationBufferMemory
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chains import ConversationChain
from groq import AsyncGroq, Groq
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

# Load environment variables from .env file
load_dotenv()

# Initialize Tavily Search - match exactly what's in the Flask app
# Initialize Tavily Search
tavily_api_key = os.getenv('TAVILY_API_KEY')
tavily_search_searcher = TavilySearch(
    max_results=5,
    include_images=True,
    include_answer=True,
    api_key=tavily_api_key
)
print("Tavily Search initialized")
print(f"Tavily search type: {type(tavily_search_searcher)}")        
print(f"Tavily search dir: {dir(tavily_search_searcher)}")

@csrf_exempt
def watson_speech_to_text(request):
    """
    Handle speech-to-text requests using IBM Watson.
    Expects audio data in the request body.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST requests are allowed'}, status=405)
    
    try:
        # Get audio data from request
        audio_data = request.body
        
        # Determine content type (default to webm)
        content_type = request.headers.get('Content-Type', 'audio/webm')
        
        # Convert speech to text
        transcription = speech_to_text(audio_data, content_type)
        
        if transcription is None:
            return JsonResponse({'error': 'Speech-to-text conversion failed'}, status=500)
        
        return JsonResponse({'transcription': transcription})
    except Exception as e:
        print(f"Error in watson_speech_to_text: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
def watson_text_to_speech(request):
    """
    Handle text-to-speech requests using IBM Watson.
    Expects JSON with 'text' and optional 'voice' parameters.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST requests are allowed'}, status=405)
    
    try:
        # Parse JSON request
        data = json.loads(request.body)
        text = data.get('text')
        voice = data.get('voice', 'en-US_AllisonV3Voice')
        
        if not text:
            return JsonResponse({'error': 'Text parameter is required'}, status=400)
        
        # Convert text to speech
        audio_data = text_to_speech(text, voice)
        
        if audio_data is None:
            return JsonResponse({'error': 'Text-to-speech conversion failed'}, status=500)
        
        # Return audio data
        response = HttpResponse(audio_data, content_type='audio/wav')
        return response
    except Exception as e:
        print(f"Error in watson_text_to_speech: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

# Dictionary to store conversation memories
conversation_memories = {}

def get_or_create_memory(session_id):
    """Get or create a conversation memory for the session"""
    if session_id not in conversation_memories:
        conversation_memories[session_id] = ConversationBufferMemory(return_messages=True)
    return conversation_memories[session_id]

def home(request):
    """View function for the home page of the site."""
    return render(request, 'personalassistant/index.html')

def signup(request):
    """View function for user registration."""
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            username = form.cleaned_data.get('username')
            messages.success(request, f'Account created for {username}! You can now log in.')
            raw_password = form.cleaned_data.get('password1')
            user = authenticate(username=username, password=raw_password)
            login(request, user)
            return redirect('personalassistant:home')
    else:
        form = UserCreationForm()
    return render(request, 'personalassistant/signup.html', {'form': form})


def login_view(request):
    """View function for user login."""
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                messages.success(request, f'Welcome back, {username}!')
                return redirect('personalassistant:dashboard')
            else:
                messages.error(request, 'Invalid username or password.')
        else:
            messages.error(request, 'Invalid username or password.')
    else:
        form = AuthenticationForm()
    return render(request, 'personalassistant/login.html', {'form': form})


def logout_view(request):
    """View function for user logout."""
    logout(request)
    messages.success(request, 'You have been logged out successfully.')
    return redirect('personalassistant:home')


@login_required
def profile(request):
    """View function for user profile."""
    return render(request, 'personalassistant/profile.html')


@login_required
def dashboard(request):
    """View function for the dashboard page."""
    return render(request, 'personalassistant/dashboard.html')


@login_required
def agento_assistant(request):
    """View function for the Agento Assistant page."""
    return render(request, 'personalassistant/agento_assistant.html')


@login_required
def ai_tutor(request):
    """View for the AI Tutor page"""
    return render(request, 'personalassistant/ai_tutor.html')

@login_required
def subject_tutor(request, subject):
    """View for the subject-specific tutor page"""
    # Map subject code to display name
    subject_names = {
        'math': 'Mathematics',
        'science': 'Science',
        'programming': 'Programming',
        'language': 'Language Arts',
        'history': 'History',
        'custom': 'Custom Subject'
    }
    
    subject_name = subject_names.get(subject, 'Subject')
    
    return render(request, 'personalassistant/subject_tutor.html', {
        'subject': subject,
        'subject_name': subject_name
    })

def process_pdf(file):
    """Extract text from a PDF file"""
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            # Write uploaded file content to temp file
            for chunk in file.chunks():
                temp_file.write(chunk)
            temp_file_path = temp_file.name
        
        # Extract text from PDF
        text = ""
        with fitz.open(temp_file_path) as pdf_document:
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text += page.get_text()
        
        # Clean up temp file
        os.unlink(temp_file_path)
        
        return text
    except Exception as e:
        print(f"Error processing PDF: {str(e)}")
        return ""

def process_docx(file):
    """Extract text from a DOCX file"""
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            # Write uploaded file content to temp file
            for chunk in file.chunks():
                temp_file.write(chunk)
            temp_file_path = temp_file.name
        
        # Extract text from DOCX
        doc = docx.Document(temp_file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Clean up temp file
        os.unlink(temp_file_path)
        
        return text
    except Exception as e:
        print(f"Error processing DOCX: {str(e)}")
        return ""

def process_txt(file):
    """Extract text from a TXT file"""
    try:
        text = ""
        for chunk in file.chunks():
            text += chunk.decode('utf-8')
        return text
    except Exception as e:
        print(f"Error processing TXT: {str(e)}")
        return ""

@login_required
@require_POST
def save_subject_context(request):
    """Handle subject context uploads (file or text)"""
    try:
        subject = request.POST.get('subject')
        context_text = request.POST.get('context_text', '')
        context_file = request.FILES.get('context_file')
        
        if not subject:
            return JsonResponse({'success': False, 'error': 'Subject is required'}, status=400)
        
        if not context_file and not context_text:
            return JsonResponse({'success': False, 'error': 'Either file or text context is required'}, status=400)
        
        # Get or create subject context for this user and subject
        subject_context, created = SubjectContext.objects.get_or_create(
            user=request.user,
            subject=subject
        )
        
        # Clear previous context data
        subject_context.context_text = ''
        if subject_context.context_file:
            subject_context.context_file.delete(save=False)
        
        # Update context text if provided
        if context_text:
            subject_context.context_text = context_text
            subject_context.file_type = None
        
        # Process file if uploaded
        if context_file:
            # Get file extension
            file_ext = context_file.name.split('.')[-1].lower()
            
            if file_ext not in ['pdf', 'docx', 'txt']:
                return JsonResponse({'success': False, 'error': 'Unsupported file format. Please upload PDF, DOCX, or TXT files only.'}, status=400)
            
            # Save file type
            subject_context.file_type = file_ext
            
            # Save file with a unique name to prevent caching issues
            timestamp = int(time.time())
            subject_context.context_file.save(
                f"{request.user.username}_{subject}_{timestamp}.{file_ext}",
                context_file
            )
            
            # Reset file pointer for processing
            context_file.seek(0)
            
            # Process file content based on type
            extracted_text = ''
            if file_ext == 'pdf':
                extracted_text = process_pdf(context_file)
            elif file_ext == 'docx':
                extracted_text = process_docx(context_file)
            elif file_ext == 'txt':
                extracted_text = process_txt(context_file)
                
            if extracted_text:
                subject_context.context_text = extracted_text
            else:
                return JsonResponse({'success': False, 'error': f'Could not extract text from the {file_ext.upper()} file. Please check the file and try again.'}, status=400)
        
        # Save changes
        subject_context.save()
        
        return JsonResponse({'success': True, 'message': 'Context updated successfully'})
    except Exception as e:
        print(f"Error in save_subject_context: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
def subject_message(request):
    """Handle subject-specific chat messages"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            message = data.get('message')
            subject = data.get('subject')
            
            if not message or not subject:
                return JsonResponse({'error': 'Message and subject are required'}, status=400)
            
            # Get session ID
            session_id = get_session_id(request)
            
            # Try to get context for this subject
            try:
                subject_context = SubjectContext.objects.get(user=request.user, subject=subject)
                context_text = subject_context.context_text or ''
            except SubjectContext.DoesNotExist:
                context_text = ''
            
            # Create a system message with the context
            if context_text:
                # Get or create memory for this session
                memory = get_or_create_memory(session_id)
                
                # Add system message with context if not already added
                if not any('SUBJECT_CONTEXT' in msg.content for msg in memory.chat_memory.messages if hasattr(msg, 'content')):
                    system_msg = f"SUBJECT_CONTEXT: {context_text}\n\nPlease use this context to help answer questions about {subject}."
                    memory.chat_memory.add_ai_message(system_msg)
            
            # Generate streaming response
            return StreamingHttpResponse(
                generate_streaming_response('meta-llama/llama-4-scout-17b-16e-instruct', message, session_id, request),
                content_type='text/event-stream'
            )
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Only POST method is allowed'}, status=405)

@login_required
def mcp_config(request):
    """View for the MCP Configuration page"""
    return render(request, 'personalassistant/mcp_config.html')


# File processing functions
def process_image(file):
    """Extract text from image using Groq's vision model"""
    try:
        # Read the image file
        image_data = file.read()
        
        # Encode the image to base64
        base64_image = base64.b64encode(image_data).decode('utf-8')
        
        # Initialize Groq client
        client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
        
        # Default query
        query = "What's in this image?"
        
        # Create chat completion with the image
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": query},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                            },
                        },
                    ],
                }
            ],
            model="llama-3.2-11b-vision-preview",
        )
        
        # Extract the response
        response = chat_completion.choices[0].message.content
        return response
    except Exception as e:
        print(f"Error processing image: {str(e)}")
        return f"Error processing image: {str(e)}"

def process_image_with_query(file, query):
    """Extract text from image using Groq's vision model with a specific query"""
    try:
        # Read the image file
        image_data = file.read()
        
        # Encode the image to base64
        base64_image = base64.b64encode(image_data).decode('utf-8')
        
        # Initialize Groq client
        client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
        
        # Use the provided query or default
        if not query or query.strip() == "":
            query = "What's in this image?"
        
        # Create chat completion with the image
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": query},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                            },
                        },
                    ],
                }
            ],
            model="llama-3.2-11b-vision-preview",
        )
        
        # Extract the response
        response = chat_completion.choices[0].message.content
        return response
    except Exception as e:
        print(f"Error processing image: {str(e)}")
        return f"Error processing image: {str(e)}"


def process_pdf(file):
    """Extract text from PDF"""
    try:
        # Create a temporary file to save the uploaded PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            file.save(temp_file.name)
            temp_file_path = temp_file.name
        
        # Open the PDF with PyMuPDF
        doc = fitz.open(temp_file_path)
        text = ""
        
        # Extract text from each page
        for page in doc:
            text += page.get_text()
        
        # Close the document and remove the temporary file
        doc.close()
        os.unlink(temp_file_path)
        
        return text
    except Exception as e:
        print(f"Error processing PDF: {str(e)}")
        return f"Error processing PDF: {str(e)}"

def process_docx(file):
    """Extract text from DOCX"""
    try:
        doc = Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error processing DOCX: {str(e)}")
        return f"Error processing DOCX: {str(e)}"

def process_txt(file):
    """Extract text from TXT"""
    try:
        text = file.read().decode('utf-8')
        return text
    except Exception as e:
        print(f"Error processing TXT: {str(e)}")
        return f"Error processing TXT: {str(e)}"

def process_csv(file):
    """Extract text from CSV"""
    try:
        df = pd.read_csv(file)
        return df.to_string()
    except Exception as e:
        print(f"Error processing CSV: {str(e)}")
        return f"Error processing CSV: {str(e)}"

def process_excel(file):
    """Extract text from Excel"""
    try:
        df = pd.read_excel(file)
        return df.to_string()
    except Exception as e:
        print(f"Error processing Excel: {str(e)}")
        return f"Error processing Excel: {str(e)}"

def process_file(file):
    """Process file based on its type"""
    filename = file.name.lower()
    
    if filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
        # Image processing is now handled separately in the upload route
        # This is to avoid double-processing
        return "Image processing handled separately"
    elif filename.endswith('.pdf'):
        return process_pdf(file)
    elif filename.endswith(('.doc', '.docx')):
        return process_docx(file)
    elif filename.endswith('.txt'):
        return process_txt(file)
    elif filename.endswith('.csv'):
        return process_csv(file)
    elif filename.endswith(('.xls', '.xlsx')):
        return process_excel(file)
    else:
        return "Unsupported file type"

def process_search_query(query):
    """Process a search query using Tavily Search API"""
    try:
        print(f"Searching with Tavily: {query}")
        
        # Use the correct method for TavilySearch
        search_results = tavily_search_searcher.invoke({"query": query})
        
        # Print the raw results for debugging
        print("Tavily search results:", json.dumps(search_results, indent=2))
        
        # Format the results for display
        formatted_results = {
            "query": query,
            "answer": search_results.get("answer", ""),
            "results": search_results.get("results", []),
            "images": search_results.get("images", []),
            "follow_up_questions": search_results.get("follow_up_questions", [])
        }
        
        return formatted_results
    except Exception as e:
        print(f"Error in Tavily search: {str(e)}")
        return {"error": str(e)}

# Get session ID from Django session
def get_session_id(request):
    """Get or create a session ID"""
    if 'session_id' not in request.session:
        request.session['session_id'] = str(uuid.uuid4())
    return request.session['session_id']

# Import Gmail utilities
from .gmail_utils import detect_email_intent, process_email_request
from .email_handler import generate_email_response

# Import Google Drive utilities
from .drive_handler import detect_drive_intent, process_drive_request, generate_drive_response

# Import user identity utilities
from .user_identity import detect_identity_update_intent, process_identity_update, get_user_identity

# API endpoints
@csrf_exempt
@login_required
def message_api(request):
    """Handle message API requests (both GET and POST)"""
    try:
        # Handle both GET and POST requests
        if request.method == 'POST':
            # For POST requests, get data from JSON body
            if request.content_type == 'application/json':
                data = json.loads(request.body)
                query = data.get('message', '')
                
                # Get session ID
                session_id = get_session_id(request)
                
                # Check if there's a file in the request
                if 'file' in data and data['file']:
                    base64_data = data['file']
                    extracted_text = process_base64_image(base64_data)
                    query = f"{query}\n\nContent extracted from image:\n{extracted_text}"
            else:
                # Handle form data if not JSON
                query = request.POST.get('message', '')
                session_id = get_session_id(request)
        else:
            # For GET requests, get data from query parameters
            query = request.GET.get('message', '')
            session_id = get_session_id(request)
        
        # Use fixed model ID as specified
        model = 'meta-llama/llama-4-scout-17b-16e-instruct'
        
        print(f"Using model: {model}")
        print(f"Query: {query}")
        
        # Check if this is an identity update request
        identity_intent = detect_identity_update_intent(query)
        
        if identity_intent and identity_intent['is_identity_update']:
            # Process identity update
            response_message = process_identity_update(request, identity_intent)
            return JsonResponse({
                'message': response_message,
                'success': True
            })
        
        # Check if this is an email request
        email_intent = detect_email_intent(query)
        print(f"Email intent detection result: {email_intent}")
        
        # Check if this is a Google Drive request
        drive_intent, drive_intent_type, drive_parameters = detect_drive_intent(query)
        print(f"Drive intent detection result: {drive_intent}, type: {drive_intent_type}, parameters: {drive_parameters}")
        
        if email_intent and email_intent['is_email_request']:
            # Get user identity information dynamically
            identity = get_user_identity(request)
            user_name = identity['name']
            user_email = identity['email']
            user_org = identity['organization']
            
            # Prepare identity information for the prompt
            identity_info = f"I am {user_name}"
            if user_org and user_org.strip():
                identity_info += f" from {user_org}"
            identity_info += ".\n\n"
            
            # Add detailed instructions for the LLM to format the email professionally
            email_prompt = (
                f"{query}\n\n"
                f"{identity_info}"
                "Please format your response as a professional email with the following structure:\n"
                "1. To: [recipient email address]\n"
                "2. Subject: [clear, concise subject line]\n"
                "3. Body: [professional email body with proper greeting, paragraphs, and closing]\n\n"
                "Guidelines for a professional email:\n"
                "- Include a proper greeting (Dear, Hello, etc.)\n"
                "- Use clear paragraphs with proper spacing\n"
                "- Include a professional closing (Sincerely, Best regards, etc.)\n"
                "- Use proper capitalization and punctuation\n"
                "- Keep the tone professional and courteous\n"
                f"- Sign the email with my name: {user_name}\n\n"
                "After the email content, add a note saying 'I'll send this email for you.'"
            )
            
            # Use a custom handler for email requests
            return StreamingHttpResponse(
                generate_email_response(model, email_prompt, query, session_id, request),
                content_type='text/event-stream'
            )
        elif drive_intent and drive_intent_type:
            print(f"Processing Drive request with intent type: {drive_intent_type}")
            # Get user identity information dynamically
            identity = get_user_identity(request)
            user_name = identity['name']
            
            # Prepare a prompt for Google Drive operations
            drive_prompt = (
                f"{query}\n\n"
                f"I am {user_name}.\n\n"
                "I'll help you with your Google Drive request. Let me process that for you."
            )
            
            # Use a custom handler for Google Drive requests
            return StreamingHttpResponse(
                generate_drive_response(model, drive_prompt, query, drive_intent_type, session_id, request, drive_parameters),
                content_type='text/event-stream'
            )
        else:
            # Regular non-email, non-drive query
            print("No specific intent detected, processing as regular query")
            return StreamingHttpResponse(
                generate_streaming_response(model, query, session_id, request),
                content_type='text/event-stream'
            )
    
    except Exception as e:
        print(f"Error in /api/message: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@login_required
def upload_file(request):
    """Handle file upload requests"""
    try:
        if 'file' not in request.FILES:
            return JsonResponse({'error': 'No file part'}, status=400)
        
        file = request.FILES['file']
        
        if file.name == '':
            return JsonResponse({'error': 'No selected file'}, status=400)
        
        # Get the query if provided
        query = request.POST.get('message', '')
        
        # Process the file based on its type
        filename = file.name.lower()
        
        if filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            # For images, use the vision model with query
            if query and query.strip():
                # Reset file pointer to beginning
                file.seek(0)
                extracted_text = process_image_with_query(file, query)
            else:
                # Reset file pointer to beginning
                file.seek(0)
                # Use default "What's in this image?" query
                extracted_text = process_image(file)
        else:
            # For other file types, use the regular processing
            extracted_text = process_file(file)
        
        return JsonResponse({'text': extracted_text})
    
    except Exception as e:
        print(f"Error in /api/upload: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@login_required
def upload_mcp_config(request):
    """Handle Google API credentials file uploads"""
    try:
        if request.method != 'POST':
            return JsonResponse({'error': 'Only POST method is allowed'}, status=405)
            
        if 'config_file' not in request.FILES:
            return JsonResponse({'error': 'No configuration file provided'}, status=400)
            
        # Get the uploaded file
        config_file = request.FILES['config_file']
        
        # Validate file type (should be JSON)
        if not config_file.name.endswith('.json'):
            return JsonResponse({'error': 'Only JSON files are allowed'}, status=400)
            
        # Create credentials directory if it doesn't exist
        credentials_dir = os.path.join('credentials')
        os.makedirs(credentials_dir, exist_ok=True)
            
        # Save the file
        file_path = os.path.join(credentials_dir, 'google_credentials.json')
        with open(file_path, 'wb+') as destination:
            for chunk in config_file.chunks():
                destination.write(chunk)
        
        # Return success response
        return JsonResponse({
            'success': True,
            'message': 'Successfully configured Google Apps integration',
            'service': 'google'
        })
        
    except Exception as e:
        print(f"Error in upload_mcp_config: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@login_required
def save_github_token(request):
    """Handle GitHub Personal Access Token saving to .env file"""
    try:
        if request.method != 'POST':
            return JsonResponse({'error': 'Only POST method is allowed'}, status=405)
            
        github_token = request.POST.get('github_token')
        if not github_token:
            return JsonResponse({'error': 'No GitHub token provided'}, status=400)
        
        # Load existing .env file if it exists
        env_file = '.env'
        env_vars = {}
        
        if os.path.exists(env_file):
            with open(env_file, 'r') as f:
                for line in f:
                    if line.strip() and not line.startswith('#'):
                        key, value = line.strip().split('=', 1)
                        env_vars[key] = value
        
        # Update or add GitHub token
        env_vars['GITHUB_TOKEN'] = github_token
        
        # Write back to .env file
        with open(env_file, 'w') as f:
            for key, value in env_vars.items():
                f.write(f"{key}={value}\n")
        
        # Return success response
        return JsonResponse({
            'success': True,
            'message': 'Successfully saved GitHub token',
            'service': 'github'
        })
        
    except Exception as e:
        print(f"Error in save_github_token: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@login_required
def disconnect_mcp(request):
    """Handle disconnection of MCP services"""
    try:
        if request.method != 'POST':
            return JsonResponse({'error': 'Only POST method is allowed'}, status=405)
            
        service = request.GET.get('service')
        if not service:
            return JsonResponse({'error': 'No service specified'}, status=400)
        
        if service == 'google':
            # Remove Google credentials file
            credentials_file = os.path.join('credentials', 'google_credentials.json')
            if os.path.exists(credentials_file):
                os.remove(credentials_file)
                
        elif service == 'github':
            # Remove GitHub token from .env file
            env_file = '.env'
            if os.path.exists(env_file):
                env_vars = {}
                with open(env_file, 'r') as f:
                    for line in f:
                        if line.strip() and not line.startswith('#'):
                            key, value = line.strip().split('=', 1)
                            if key != 'GITHUB_TOKEN':
                                env_vars[key] = value
                
                with open(env_file, 'w') as f:
                    for key, value in env_vars.items():
                        f.write(f"{key}={value}\n")
        else:
            return JsonResponse({'error': f'Unknown service: {service}'}, status=400)
        
        # Return success response
        return JsonResponse({
            'success': True,
            'message': f'Successfully disconnected {service}',
            'service': service
        })
        
    except Exception as e:
        print(f"Error in disconnect_mcp: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def check_mcp_status(request):
    """Check connection status of MCP services"""
    try:
        # Check Google connection status
        google_connected = os.path.exists(os.path.join('credentials', 'google_credentials.json'))
        
        # Check GitHub connection status
        github_connected = False
        env_file = '.env'
        if os.path.exists(env_file):
            with open(env_file, 'r') as f:
                for line in f:
                    if line.strip().startswith('GITHUB_TOKEN='):
                        token = line.strip().split('=', 1)[1]
                        if token and token.strip():
                            github_connected = True
                            break
        
        # Return status
        return JsonResponse({
            'success': True,
            'google_connected': google_connected,
            'github_connected': github_connected
        })
        
    except Exception as e:
        print(f"Error in check_mcp_status: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)


def tavily_search(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            query = data.get('query', '')
            
            if not query:
                return JsonResponse({'error': 'Query is required'}, status=400)
            
            # Get session ID for memory management
            session_id = request.session.session_key
            if not session_id:
                request.session.create()
                session_id = request.session.session_key
            
            search_results = process_search_query(query)
            
            # Add search response to memory buffer
            try:
                if session_id in conversation_memories:  # Use conversation_memories instead of conversation_memory
                    memory = conversation_memories[session_id]
                    
                    # Add user query to memory
                    memory.chat_memory.add_user_message(query)
                    
                    # Create a formatted response message from search results
                    search_response = ""
                    if search_results.get('answer'):
                        search_response = f"Search results: {search_results['answer']}"
                    elif search_results.get('results') and len(search_results['results']) > 0:
                        search_response = "Search results:\n"
                        for idx, result in enumerate(search_results['results'][:3]):
                            search_response += f"{idx+1}. {result['title']}: {result['content'][:100]}...\n"
                    else:
                        search_response = "No relevant search results found."
                    
                    # Add search response to memory
                    memory.chat_memory.add_ai_message(search_response)
            except Exception as mem_error:
                print(f"Memory error (non-critical): {str(mem_error)}")  # Log but don't fail the request
            
            return JsonResponse(search_results)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Only POST method is allowed'}, status=405)

def generate_streaming_response(model, query, session_id, request):
    """Generate a streaming response from the model with conversation memory"""
    try:
        print(f"Generating streaming response with model: {model}")
        print(f"Query: {query}")
        
        # Get or create memory for this session
        memory = get_or_create_memory(session_id)
        
        # Add the human message to memory
        memory.chat_memory.add_user_message(query)
        
        # Get the conversation history
        messages = memory.chat_memory.messages
        
        # Convert to the format expected by the Groq API
        groq_messages = []
        
        # Add system message
        groq_messages.append({
            "role": "system",
            "content": "You are a helpful assistant specialized in coding and study-related responses."
        })
        
        # Add conversation history
        for msg in messages:
            if isinstance(msg, HumanMessage):
                groq_messages.append({"role": "user", "content": msg.content})
            elif isinstance(msg, AIMessage):
                groq_messages.append({"role": "assistant", "content": msg.content})
        
        # Initialize Groq client
        api_key = os.getenv('GROQ_API_KEY')
        client = Groq(api_key=api_key)
        
        # Create the streaming completion using the direct Groq client
        stream = client.chat.completions.create(
            messages=groq_messages,
            model='meta-llama/llama-4-scout-17b-16e-instruct',  # Use fixed model ID
            temperature=0.5,
            max_tokens=4000,
            top_p=1,
            stop=None,
            stream=True,
        )
        
        # Yield each chunk as a server-sent event
        yield f"data: {json.dumps({'status': 'start'})}\n\n"
        
        full_response = ""
        for chunk in stream:
            if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                content = chunk.choices[0].delta.content
                full_response += content
                
                # Yield the chunk as a server-sent event
                yield f"data: {json.dumps({'content': content})}\n\n"
        
        # Add the full response to memory
        memory.chat_memory.add_ai_message(full_response)
        
        # Signal the end of the stream
        yield f"data: {json.dumps({'status': 'done'})}\n\n"
        
    except Exception as e:
        print(f"Error in generate_streaming_response: {str(e)}")
        yield f"data: {json.dumps({'status': 'error', 'message': str(e)})}\n\n"


@csrf_exempt
def watson_speech_to_text(request):
    """
    API endpoint to convert speech to text using IBM Watson Speech-to-Text service.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST method is allowed'}, status=405)
    
    try:
        # Import the Watson services module
        from .watson_services import speech_to_text
        
        # Get audio data from request
        audio_data = request.body
        content_type = request.headers.get('Content-Type', 'audio/webm')
        
        # Convert speech to text
        transcription = speech_to_text(audio_data, content_type)
        
        if transcription is not None:
            return JsonResponse({'transcription': transcription})
        else:
            return JsonResponse({'error': 'Failed to transcribe audio'}, status=500)
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
def watson_text_to_speech(request):
    """
    API endpoint to convert text to speech using IBM Watson Text-to-Speech service.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST method is allowed'}, status=405)
    
    try:
        # Import the Watson services module
        from .watson_services import text_to_speech
        
        # Get text from request
        data = json.loads(request.body)
        text = data.get('text', '')
        voice = data.get('voice', 'en-US_AllisonV3Voice')
        
        if not text:
            return JsonResponse({'error': 'Text is required'}, status=400)
        
        # Convert text to speech
        audio_data = text_to_speech(text, voice)
        
        if audio_data is not None:
            # Return audio data as response
            response = HttpResponse(audio_data, content_type='audio/wav')
            response['Content-Disposition'] = 'attachment; filename="speech.wav"'
            return response
        else:
            return JsonResponse({'error': 'Failed to synthesize speech'}, status=500)
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
